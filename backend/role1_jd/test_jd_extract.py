#!/usr/bin/env python3
"""
Standalone end-to-end test for the JD Analytics endpoint.

Usage (from the backend/ directory, with venv activated):
    python -m role1_jd.test_jd_extract

What it does:
  1. Creates a sample Job Description as a temporary DOCX file.
  2. Starts the FastAPI app on a free port using uvicorn in-process.
  3. Sends a POST request to /api/jd/extract with the sample DOCX.
  4. Prints the returned JSON.
  5. Validates the response against the shared ExtractedSkillList schema.
  6. Reports PASS / FAIL.
"""

import json
import sys
import os
import tempfile
import threading
import time

import requests
import uvicorn
from docx import Document

# ──────────────────────────────────────────────
# 1. Create a realistic sample JD as DOCX
# ──────────────────────────────────────────────

SAMPLE_JD_TEXT = """
Senior Backend Engineer – Cloud Platform

About the Company
TechNova Solutions is a fast-growing SaaS company building next-generation
developer tools powered by AI.

Key Responsibilities
- Design and implement scalable RESTful APIs using Python and FastAPI.
- Build and manage cloud infrastructure on AWS (EC2, Lambda, S3, RDS).
- Write efficient SQL queries and optimise PostgreSQL databases.
- Implement CI/CD pipelines using GitHub Actions and Docker.
- Collaborate with front-end engineers to define API contracts.
- Participate in system design reviews and architecture discussions.
- Mentor junior engineers through code reviews and pair programming.

What We're Looking For
- 4+ years of professional software engineering experience.
- Strong proficiency in Python and object-oriented design.
- Deep understanding of data structures and algorithms.
- Experience with cloud platforms (AWS preferred).
- Solid knowledge of relational databases and SQL.
- Excellent communication and teamwork skills.
- Familiarity with containerisation (Docker, Kubernetes).
- Experience with microservices architecture and system design.
- Knowledge of networking fundamentals (TCP/IP, HTTP, DNS).
- Understanding of operating system concepts (Linux).

Nice to Have
- Experience with machine learning or AI/ML pipelines.
- Contributions to open-source projects.
- Experience with real-time data processing.

Benefits
- Competitive salary and equity.
- Flexible remote work policy.
- Learning and development budget.
"""


def create_sample_docx() -> str:
    """Create a temporary DOCX file with the sample JD and return its path."""
    doc = Document()
    doc.add_heading("Senior Backend Engineer – Cloud Platform", level=1)
    for paragraph in SAMPLE_JD_TEXT.strip().split("\n"):
        doc.add_paragraph(paragraph)

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False, prefix="sample_jd_")
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


# ──────────────────────────────────────────────
# 2. Start FastAPI server in a background thread
# ──────────────────────────────────────────────

TEST_PORT = 8899
BASE_URL = f"http://127.0.0.1:{TEST_PORT}"


def start_server():
    """Run the uvicorn server (blocking)."""
    uvicorn.run("main:app", host="127.0.0.1", port=TEST_PORT, log_level="warning")


def wait_for_server(timeout: int = 15):
    """Poll the health endpoint until the server is ready."""
    deadline = time.time() + timeout
    while time.time() < deadline:
        try:
            r = requests.get(f"{BASE_URL}/health", timeout=2)
            if r.status_code == 200:
                return True
        except requests.ConnectionError:
            pass
        time.sleep(0.5)
    return False


# ──────────────────────────────────────────────
# 3. Main test logic
# ──────────────────────────────────────────────

def main():
    print("=" * 60)
    print("  JD Analytics – End-to-End Test")
    print("=" * 60)

    # Check for API key
    from dotenv import load_dotenv
    load_dotenv()
    if not os.environ.get("GEMINI_API_KEY"):
        print("\n❌  GEMINI_API_KEY not set. Create a .env file with your key.")
        print("   Example:  GEMINI_API_KEY=your_key_here")
        sys.exit(1)

    # Create sample DOCX
    docx_path = create_sample_docx()
    print(f"\n📄  Created sample JD: {docx_path}")

    # Start server
    print(f"\n🚀  Starting FastAPI server on port {TEST_PORT}...")
    server_thread = threading.Thread(target=start_server, daemon=True)
    server_thread.start()

    if not wait_for_server():
        print("❌  Server failed to start within 15 seconds.")
        sys.exit(1)
    print("✅  Server is ready.")

    # Send POST request
    print("\n📤  Sending POST /api/jd/extract ...")
    try:
        with open(docx_path, "rb") as f:
            response = requests.post(
                f"{BASE_URL}/api/jd/extract",
                files={"file": ("sample_jd.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                params={"company": "TechNova Solutions", "role": "Senior Backend Engineer"},
                timeout=60,
            )
    finally:
        os.unlink(docx_path)  # Clean up temp file

    # Check response
    print(f"\n📥  Response Status: {response.status_code}")

    if response.status_code != 200:
        print(f"❌  Request failed!\n{response.text}")
        sys.exit(1)

    result = response.json()
    print(f"\n📋  Extracted JSON ({len(result.get('skills', []))} skills):\n")
    print(json.dumps(result, indent=2))

    # Validate against shared schema
    print("\n🔍  Validating against ExtractedSkillList schema...")
    try:
        # Add the parent dir so we can import shared
        sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        from shared.schemas import ExtractedSkillList
        validated = ExtractedSkillList(**result)
        print(f"✅  Validation PASSED — {len(validated.skills)} skills extracted.")
        print(f"    source_type : {validated.source_type}")
        print(f"    source_file : {validated.source_file}")
        print(f"    company     : {validated.company}")
        print(f"    role        : {validated.role}")
        print("\n    Skills breakdown:")
        for skill in validated.skills:
            print(f"      • [{skill.category_code}] {skill.skill_name} "
                  f"(confidence: {skill.confidence})")
    except Exception as exc:
        print(f"❌  Validation FAILED: {exc}")
        sys.exit(1)

    print("\n" + "=" * 60)
    print("  ✅  ALL TESTS PASSED")
    print("=" * 60)


if __name__ == "__main__":
    main()
